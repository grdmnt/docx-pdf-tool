from docx import Document
from docx.shared import Inches
import os
import csv

FONT = 'Lato Medium'
BOLD = True
INDENT_INCHES = 1.5

DOCX_TEMPLATE = 'header-probi.docx'
CSV_FILE = 'probi_members.csv'

def replace_variable(paragraph, variable, value):
    return paragraph.replace(variable, value)

def get_head_substring(paragraph, variable):
    return paragraph[:paragraph.index(variable)]

def get_tail_substring(paragraph, variable):
    return paragraph[paragraph.index(variable) + len(variable):]

def convert_docx_to_pdf(filename):
    os.system('unoconv -f pdf ' + filename)

def move_file_to_generated_files(filename):
    os.system('mv ' + filename + ' generated_files')

def create_folder_for_generated_files():
    os.system('mkdir generated_files')

def read_csv(filename):
    variables = []
    reader = csv.reader(open(filename))

    content = []
    for row in reader:
        if not variables:
            variables = row
            continue
        content.append(row)
    return variables, content

def change_content(paragraph, string, variables, values, variable, clear = False):
    head_string = get_head_substring(string, variable)
    tail_string = get_tail_substring(string, variable)

    if not clear:
        paragraph.clear()

    p = paragraph

    p.paragraph_format.left_indent = Inches(INDENT_INCHES)

    run = p.add_run(head_string)
    font = run.font
    font.name = FONT

    run = p.add_run(values[variables.index(variable)])
    run.bold = True
    font = run.font
    font.name = FONT

    if '{{' in tail_string:
        head = tail_string.index('{{')
        tail = tail_string.index('}}') + 2
        new_variable = tail_string[head:tail]
        change_content(paragraph, tail_string, variables, values, new_variable, True)
        return
    run = p.add_run(tail_string)
    font = run.font
    font.name = FONT

def create_docx(variables, values, output_file):
    in_file = DOCX_TEMPLATE
    processed_file = output_file

    document = Document(in_file)

    for paragraph in document.paragraphs:
        for variable in variables:
            if variable in paragraph.text:
                change_content(paragraph, paragraph.text, variables, values, variable)


    document.save(processed_file)

    convert_docx_to_pdf(processed_file)


def main():
    variables, content = read_csv(CSV_FILE)

    if not os.path.isdir('generated_files'):
        create_folder_for_generated_files()

    output_filename_index = variables.index('{{FILENAME}}')

    for entry in content:
        create_docx(variables, entry, entry[output_filename_index]+'.docx')
        move_file_to_generated_files(entry[output_filename_index]+'.docx')
        move_file_to_generated_files(entry[output_filename_index]+'.pdf')

if __name__ == "__main__":
    main()

